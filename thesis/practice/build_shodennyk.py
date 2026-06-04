#!/usr/bin/env python3
"""Build the practice diary (shodennyk.docx) by filling the official template.

The щоденник is a fixed official Word form. To match the institutional layout
1:1, we do NOT rebuild it — we open the official template
`shodennyk-template.docx` (the Odesa Polytechnic diary form) and substitute only
the data fields from `shodennyk-data.md`, preserving every table, border, font
and page setting byte-for-byte.

Data filled: student name, course/group, enterprise, university supervisor,
the Календарний графік works + week marks, the Робочі записи entries, and the
Індивідуальне завдання. Signature / seal / grade fields and the enterprise
review stay blank (filled by hand).

Usage:  python build_shodennyk.py
"""
import re
import sys
from pathlib import Path

from docx import Document

HERE = Path(__file__).parent
TEMPLATE = HERE / "shodennyk-template.docx"
DATA = HERE / "shodennyk-data.md"

# Identity tokens present in the template (NeoFauna sample) that must be
# replaced with the practitioner's data. The template's practice dates
# (02–27 лютого 2026) already match shodennyk-data.md, so dates are untouched.
TPL_GROUP = "АС-223"
TPL_ENTERPRISE = "ФОП Мала Валерія Олександрівна"
TPL_SUPERVISOR = "Зіноватна С.Л."
# Fixed-position template tables (stable for this official form):
T_REVIEW = 6        # «Відгук і оцінка роботи здобувача» — enterprise review
T_WORK_A = 5        # «Робочі записи» — first block (5 one-col rows)
T_WORK_B = 4        # «Робочі записи» — overflow block (clear it)
T_CALENDAR = 8      # «Календарний графік» — №/Назви робіт/тижні 1–4/відмітки


# --------------------------------------------------------------------------- #
# Data                                                                        #
# --------------------------------------------------------------------------- #
def _load_data(path: Path) -> dict:
    raw = path.read_text(encoding="utf-8")
    m = re.search(r"```ya?ml\s*(.*?)```", raw, re.DOTALL)
    payload = m.group(1) if m else raw
    try:
        import yaml
    except ImportError:
        sys.exit("PyYAML required: pip install pyyaml (into thesis/.venv)")
    data = yaml.safe_load(payload)
    if not isinstance(data, dict):
        sys.exit(f"{path.name}: expected a YAML mapping of form fields")
    return data


def _txt(v) -> str:
    return "" if v is None else str(v)


def _ddmm(iso: str) -> str:
    """'2026-02-02' -> '02.02'; pass through anything not ISO."""
    m = re.match(r"\d{4}-(\d{2})-(\d{2})", _txt(iso))
    return f"{m.group(2)}.{m.group(1)}" if m else _txt(iso)


def _supervisor_short(full: str) -> str:
    """'Надія БЕГЛОВА' -> 'Беглова Н.' (surname + first initial), to match the
    template's 'Прізвище І.' convention. Pass through if already short."""
    parts = _txt(full).split()
    if len(parts) < 2:
        return _txt(full)
    first, surname = parts[0], parts[-1]
    return f"{surname.title()} {first[0]}."


# --------------------------------------------------------------------------- #
# Run-safe text editing (preserves formatting)                                #
# --------------------------------------------------------------------------- #
def _iter_paragraphs(doc):
    # No merged-cell dedup: replacement is idempotent, so visiting a merged
    # cell's paragraphs more than once is harmless. (Deduping by id(c._tc) is
    # unsafe — lxml hands out fresh proxy wrappers whose ids can collide.)
    yield from doc.paragraphs
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                yield from c.paragraphs


def _set_para_text(p, text: str) -> None:
    """Write `text` into the paragraph keeping its first run's formatting."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    elif text:
        p.add_run(text)


def _replace_in_paragraph(p, old: str, new: str) -> bool:
    if old not in p.text:
        return False
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    # token spans several runs → collapse into the first run
    _set_para_text(p, p.text.replace(old, new))
    return True


def _replace_global(doc, old: str, new: str) -> int:
    return sum(_replace_in_paragraph(p, old, new) for p in _iter_paragraphs(doc))


def _set_cell(cell, text: str) -> None:
    _set_para_text(cell.paragraphs[0], text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


# --------------------------------------------------------------------------- #
# Section fillers                                                             #
# --------------------------------------------------------------------------- #
def _fill_student_name(doc, name: str) -> None:
    # The «Студент _____» line lives in the second header table, row 0,
    # immediately to the right of the «Студент» label.
    row = doc.tables[1].rows[0]
    _set_cell(row.cells[1], name)


def _fill_individual_task(doc, task: str) -> None:
    task = _txt(task).strip()
    if not task:
        return
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip().startswith("Індивідуальне завдання"):
            if i + 1 < len(paras):
                _set_para_text(paras[i + 1], task)
            else:
                p.add_run(" " + task)
            return


def _fill_calendar(doc, items: list) -> None:
    table = doc.tables[T_CALENDAR]
    # data rows start at index 2 (two header rows); template has 8 data rows.
    for i, it in enumerate(items):
        r = 2 + i
        if r >= len(table.rows):
            break
        cells = table.rows[r].cells
        weeks = set(it.get("weeks", []) or [])
        _set_cell(cells[0], _txt(it.get("n", i + 1)))
        _set_cell(cells[1], _txt(it.get("work", "")))
        for wk in (1, 2, 3, 4):
            _set_cell(cells[1 + wk], "х" if wk in weeks else "")


def _fill_work_entries(doc, entries: list) -> None:
    lines = [f'{_ddmm(e.get("from", ""))} – {_ddmm(e.get("to", ""))} '
             f'{_txt(e.get("text", "")).strip()}' for e in entries]
    ta = doc.tables[T_WORK_A]
    for i, row in enumerate(ta.rows):
        _set_cell(row.cells[0], lines[i] if i < len(lines) else "")
    # overflow block: blank it (template held leftover sample entries)
    for row in doc.tables[T_WORK_B].rows:
        _set_cell(row.cells[0], "")


def _clear_enterprise_review(doc) -> None:
    # «Відгук і оцінка...»: keep the enterprise name (row 0, replaced globally),
    # blank the prose written by the sample's enterprise supervisor.
    table = doc.tables[T_REVIEW]
    for row in table.rows[2:]:
        _set_cell(row.cells[0], "")


# --------------------------------------------------------------------------- #
# Build                                                                       #
# --------------------------------------------------------------------------- #
def build_shodennyk() -> None:
    if not TEMPLATE.exists():
        sys.exit(f"missing template: {TEMPLATE}")
    d = _load_data(DATA)
    doc = Document(str(TEMPLATE))

    n_group = _replace_global(doc, TPL_GROUP, _txt(d.get("group", TPL_GROUP)))
    n_ent = _replace_global(doc, TPL_ENTERPRISE, _txt(d.get("enterprise", "")))
    n_sup = _replace_global(doc, TPL_SUPERVISOR,
                            _supervisor_short(d.get("kerivnyk_vnz", "")))

    _fill_student_name(doc, _txt(d.get("student", "")))
    _fill_individual_task(doc, d.get("individual_task", ""))
    _fill_calendar(doc, d.get("calendar", []) or [])
    _fill_work_entries(doc, d.get("work_entries", []) or [])
    _clear_enterprise_review(doc)

    out = HERE / "shodennyk.docx"
    doc.save(str(out))
    print(f"OK -> {out}  (group×{n_group}, enterprise×{n_ent}, supervisor×{n_sup})")


def main() -> None:
    build_shodennyk()


if __name__ == "__main__":
    main()
